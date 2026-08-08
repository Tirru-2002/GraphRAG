"""
Graph backend code with spaCy NER instead of LLM
Place this file next to app.py and import GraphRAGChatbot from it.
"""
import os
import logging
from pathlib import Path
from typing import List, Optional, Dict

import nest_asyncio
from dotenv import load_dotenv
import pypdf
import docx
import csv
from neo4j import GraphDatabase
import spacy
from spacy.tokens import Doc

# --- Docling (optional) ---------------------------------------------------
# Docling parses PDF / DOCX / PPTX / HTML / images locally.
# No API key needed and no internet needed after the model weights are
# cached the first time. If it is not installed, the app still works,
# it just falls back to the old pypdf / python-docx extractors below.
try:
    from docling.document_converter import DocumentConverter
    from docling.chunking import HybridChunker
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
# ---------------------------------------------------------------------------

# allow nested event loops when using asyncio.run inside Streamlit
nest_asyncio.apply()

# load .env if present
load_dotenv()

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process multiple document formats into plain text."""

    # Docling loads its layout/OCR models once. We cache the converter
    # here so every call does not reload the models from scratch.
    _docling_converter = None

    # ---------------- Docling functions ----------------

    @staticmethod
    def _get_docling_converter():
        """Create the Docling converter once, then reuse it (lazy singleton)."""
        if DocumentProcessor._docling_converter is None:
            DocumentProcessor._docling_converter = DocumentConverter()
        return DocumentProcessor._docling_converter

    @staticmethod
    def extract_with_docling(file_path: str) -> str:
        """
        Extract a document's content as Markdown using Docling.

        Docling supports PDF, DOCX, PPTX, HTML, images (PNG/JPEG) and more.
        It keeps headings, lists and table structure, which plain
        pypdf/python-docx text extraction loses. This runs 100% locally
        on CPU - no CUDA and no API key required.
        """
        if not DOCLING_AVAILABLE:
            raise ImportError(
                "Docling is not installed. Run: pip install docling"
            )

        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            converter = DocumentProcessor._get_docling_converter()
            result = converter.convert(str(file_path))
            return result.document.export_to_markdown()
        except Exception as e:
            logger.error(f"Error processing file with Docling: {e}")
            raise

    @staticmethod
    def extract_with_docling_chunks(file_path: str) -> List[str]:
        """
        Extract a document with Docling AND split it into RAG-ready chunks
        using Docling's HybridChunker.

        This is the recommended function to call before feeding text into
        SpacyGraphExtractor / Neo4j below - each returned chunk is already
        a clean, self-contained piece of text, sized for embedding/graph
        extraction, instead of one giant blob of text.
        """
        if not DOCLING_AVAILABLE:
            raise ImportError(
                "Docling is not installed. Run: pip install docling"
            )

        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            converter = DocumentProcessor._get_docling_converter()
            result = converter.convert(str(file_path))

            chunker = HybridChunker()
            chunks = list(chunker.chunk(result.document))

            return [chunk.text for chunk in chunks]
        except Exception as e:
            logger.error(f"Error chunking file with Docling: {e}")
            raise

    # ---------------- Original fallback extractors (unchanged) ----------------

    @staticmethod
    def extract_pdf(file_path: str) -> str:
        try:
            text = ""
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = pypdf.PdfReader(pdf_file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page.extract_text() or ""
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
            return text
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise

    @staticmethod
    def extract_docx(file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise

    @staticmethod
    def extract_csv(file_path: str) -> str:
        try:
            text = ""
            with open(file_path, 'r', encoding='utf-8') as csv_file:
                reader = csv.DictReader(csv_file)
                for row in reader:
                    text += str(row) + "\n"
            return text
        except Exception as e:
            logger.error(f"Error processing CSV: {e}")
            raise

    @staticmethod
    def extract_txt(file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as txt_file:
                text = txt_file.read()
            return text
        except Exception as e:
            logger.error(f"Error processing TXT: {e}")
            raise

    # ---------------- Main entry point ----------------

    @staticmethod
    def process_document(file_path: str, use_docling: bool = True) -> str:
        """
        Convert a document into plain/markdown text.

        use_docling=True (default):
            Use Docling for PDF / DOCX / PPTX / HTML / images, since it
            gives cleaner text and keeps table structure. If Docling is
            not installed, or the conversion fails for any reason, this
            automatically falls back to the old pypdf / python-docx code.
        use_docling=False:
            Always use the old pypdf / python-docx extractors only.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()
        docling_supported = {'.pdf', '.docx', '.pptx', '.html', '.htm',
                              '.png', '.jpg', '.jpeg'}

        if use_docling and DOCLING_AVAILABLE and extension in docling_supported:
            try:
                return DocumentProcessor.extract_with_docling(str(file_path))
            except Exception as e:
                logger.warning(f"Docling extraction failed, falling back to old extractor: {e}")

        processors = {
            '.pdf': DocumentProcessor.extract_pdf,
            '.docx': DocumentProcessor.extract_docx,
            '.csv': DocumentProcessor.extract_csv,
            '.txt': DocumentProcessor.extract_txt
        }

        if extension not in processors:
            raise ValueError(f"Unsupported file format: {extension}")

        return processors[extension](str(file_path))


class SpacyGraphExtractor:
    """Extract knowledge graph using spaCy NER and dependency parsing."""

    def __init__(self, model_name: str = "en_core_web_sm"):
        """Initialize spaCy model."""
        try:
            self.nlp = spacy.load(model_name)
            logger.info(f"Loaded spaCy model: {model_name}")
        except OSError:
            logger.warning(f"Model {model_name} not found. Downloading...")
            os.system(f"python -m spacy download {model_name}")
            self.nlp = spacy.load(model_name)

    def extract_entities(self, text: str) -> List[Dict]:
        """Extract named entities from text."""
        doc = self.nlp(text)
        entities = []

        for ent in doc.ents:
            entities.append({
                'id': ent.text,
                'type': ent.label_,
                'text': ent.text
            })

        # Remove duplicates based on text
        unique_entities = []
        seen = set()
        for entity in entities:
            if entity['text'] not in seen:
                seen.add(entity['text'])
                unique_entities.append(entity)

        return unique_entities

    def extract_relationships(self, text: str) -> List[Dict]:
        """Extract relationships using dependency parsing."""
        doc = self.nlp(text)
        relationships = []

        for token in doc:
            # Find verb-based relationships
            if token.pos_ == "VERB":
                subject = None
                obj = None

                # Find subject and object
                for child in token.children:
                    if child.dep_ in ["nsubj", "nsubjpass"]:  # Nominal Subject
                        subject = child
                    elif child.dep_ in ["dobj", "pobj", "attr"]:  # Direct Object
                        obj = child

                # Create relationship if both subject and object found
                if subject and obj:
                    # Get entity types
                    subject_type = self._get_entity_type(subject, doc)
                    obj_type = self._get_entity_type(obj, doc)

                    relationships.append({
                        'source': {
                            'id': subject.text,
                            'type': subject_type
                        },
                        'target': {
                            'id': obj.text,
                            'type': obj_type
                        },
                        'type': self._classify_relationship(token.text),  # every text here is verb, because token.pos_ == "VERB"
                        'verb': token.text
                    })

        return relationships

    def _get_entity_type(self, token, doc) -> str:
        """Get entity type for a token."""
        # Check if token is part of a named entity
        for ent in doc.ents:
            if token.i >= ent.start and token.i < ent.end:
                return ent.label_

        # Default type based on POS tag
        if token.pos_ == "PROPN":
            return "ENTITY"
        elif token.pos_ == "NOUN":
            return "CONCEPT"
        else:
            return "UNKNOWN"

    def _classify_relationship(self, verb: str) -> str:
        """Classify relationship type based on verb."""
        verb_lower = verb.lower()

        # Define relationship mappings
        work_verbs = ["works", "work", "worked", "working", "employed"]
        manage_verbs = ["manages", "manage", "managed", "managing", "leads", "lead"]
        create_verbs = ["creates", "create", "created", "builds", "built", "develops"]
        owns_verbs = ["owns", "own", "owned", "possesses", "has"]
        located_verbs = ["located", "lives", "resides", "based"]

        if verb_lower in work_verbs:
            return "WORKS_FOR"
        elif verb_lower in manage_verbs:
            return "MANAGES"
        elif verb_lower in create_verbs:
            return "CREATES"
        elif verb_lower in owns_verbs:
            return "OWNS"
        elif verb_lower in located_verbs:
            return "LOCATED_IN"
        else:
            return "RELATED_TO"

    def extract_graph(self, text: str) -> Dict:
        """Extract complete graph from text."""
        entities = self.extract_entities(text)
        relationships = self.extract_relationships(text)

        return {
            'nodes': entities,
            'relationships': relationships
        }


class Neo4jGraphDB:
    """Manage Neo4j graph database operations."""

    def __init__(self, uri: str, user: str, password: str):
        try:
            self.driver = GraphDatabase.driver(uri, auth=(user, password))
            self.driver.verify_connectivity()
            logger.info("Connected to Neo4j successfully")
        except Exception as e:
            logger.error(f"Failed to connect to Neo4j: {e}")
            self.driver = None

    def is_connected(self) -> bool:
        return self.driver is not None

    def close(self):
        if self.driver:
            self.driver.close()

    def clear_database(self):
        if not self.is_connected():
            return
        try:
            with self.driver.session() as session:
                session.run("MATCH (n) DETACH DELETE n")
            logger.info("Database cleared")
        except Exception as e:
            logger.error(f"Error clearing database: {e}")

    def upload_graph_data_from_spacy(self, graph_data: Dict):
        """Upload graph data extracted by spaCy.

        - Sanitize labels and relationship types.
        - Create uniqueness constraints for each label (IF NOT EXISTS).
        - MERGE nodes by label+id.
        - For relationships: attempt label-qualified MERGE first; then do an id-only MERGE fallback.
        - At end, run verification queries and log counts by relationship type.
        """
        import re
        from collections import defaultdict

        def _norm_id(value: str) -> str:
            return str(value).strip().lower()

        def _sanitize_label(label: str) -> str:
            if not label or not isinstance(label, str):
                label = "ENTITY"
            lbl = re.sub(r'[^0-9A-Za-z_]', '_', label.strip())
            if re.match(r'^\d', lbl):
                lbl = "Num_" + lbl
            return lbl or "ENTITY"

        def _sanitize_rel_type(rel: str) -> str:
            if not rel or not isinstance(rel, str):
                rel = "RELATED_TO"
            rt = re.sub(r'[^0-9A-Za-z_]', '_', rel.strip()).upper()
            return rt or "RELATED_TO"

        if not self.is_connected():
            logger.warning("Not connected to Neo4j")
            return

        if not graph_data:
            logger.warning("No graph data to upload")
            return

        total_nodes = 0
        total_relationships = 0

        try:
            with self.driver.session() as session:
                # --- Nodes: group by sanitized label
                nodes = graph_data.get('nodes', []) or []
                node_types_map = {}
                for node in nodes:
                    raw_type = node.get('type', 'ENTITY')
                    label = _sanitize_label(raw_type)
                    node_types_map.setdefault(label, []).append(node)

                # Create uniqueness constraint for each label (id)
                for label in node_types_map.keys():
                    try:
                        constraint_query = (
                            f"CREATE CONSTRAINT IF NOT EXISTS FOR (n:`{label}`) "
                            "REQUIRE n.id IS UNIQUE"
                        )
                        session.run(constraint_query)
                    except Exception as e:
                        logger.warning(f"Could not create constraint for {label}: {e}")

                # MERGE nodes by label + id
                for label, typed_nodes in node_types_map.items():
                    prepared = []
                    for node in typed_nodes:
                        node_id = _norm_id(node.get('id') or node.get('text') or node.get('name') or "")

                        if not node_id:
                            continue
                        node_text = (node.get('text') or node.get('id') or node.get('name') or "").strip()
                        prepared.append({"id": node_id, "name": node_id, "text": node_text})
                    if not prepared:
                        continue
                    node_query = f"""
                    UNWIND $nodes AS node
                    MERGE (n:`{label}` {{id: node.id}})
                    SET n.name = node.name, n.text = node.text
                    """
                    session.run(node_query, nodes=prepared)
                    total_nodes += len(prepared)
                    logger.info(f"Created/merged {len(prepared)} nodes of type {label}")

                # --- Relationships: prepare and create with fallback
                relationships = graph_data.get('relationships', []) or []
                if relationships:
                    rel_type_map = defaultdict(list)
                    for rel in relationships:
                        rtype_raw = rel.get('type') or rel.get('verb') or "RELATED_TO"
                        rtype = _sanitize_rel_type(rtype_raw)

                        source = rel.get('source') or {}
                        target = rel.get('target') or {}
                        src_id = _norm_id(source.get('id') or source.get('text') or source.get('name') or "")
                        tgt_id = _norm_id(target.get('id') or target.get('text') or target.get('name') or "")

                        if not src_id or not tgt_id:
                            logger.debug(f"Skipping relationship with missing ids: {rel}")
                            continue

                        src_label = _sanitize_label(source.get('type') or "ENTITY")
                        tgt_label = _sanitize_label(target.get('type') or "ENTITY")

                        rel_type_map[rtype].append({
                            "src_id": src_id,
                            "tgt_id": tgt_id,
                            "src_label": src_label,
                            "tgt_label": tgt_label,
                            "verb": rel.get('verb') or ""
                        })

                    for rtype, rels in rel_type_map.items():
                        # group by (src_label, tgt_label)
                        pair_map = defaultdict(list)
                        for r in rels:
                            pair_map[(r['src_label'], r['tgt_label'])].append(r)

                        for (src_lbl, tgt_lbl), pair_rels in pair_map.items():
                            prepared = [{"src_id": r["src_id"], "tgt_id": r["tgt_id"], "verb": r.get("verb", "")} for r in pair_rels]
                            if not prepared:
                                continue

                            # Label-qualified MATCH + MERGE relationship (safer than unlabeled) --> best practice
                            try:
                                fallback_query = f"""
                                UNWIND $pairs AS rel
                                MATCH (source:`{src_lbl}` {{id: rel.src_id}})
                                MATCH (target:`{tgt_lbl}` {{id: rel.tgt_id}})
                                MERGE (source)-[rr:`{rtype}`]->(target)
                                SET rr.verb = rel.verb
                                """
                                session.run(fallback_query, pairs=prepared)
                                total_relationships += len(prepared)
                                logger.info(f"Fallback (label-MATCH) created {len(prepared)} relationships of type {rtype} between {src_lbl} -> {tgt_lbl}")
                                continue
                            except Exception as e:
                                logger.warning(f"Fallback label-MATCH creation failed for {rtype} {src_lbl}->{tgt_lbl}: {e}")

                            # Last resort (id-only) if the above fail
                            try:
                                id_only_query = f"""
                                UNWIND $pairs AS rel
                                MATCH (source {{id: rel.src_id}})
                                MATCH (target {{id: rel.tgt_id}})
                                MERGE (source)-[rr:`{rtype}`]->(target)
                                SET rr.verb = rel.verb
                                """
                                session.run(id_only_query, pairs=prepared)
                                total_relationships += len(prepared)
                                logger.info(f"Final id-only create {len(prepared)} of type {rtype} between {src_lbl}->{tgt_lbl}")
                            except Exception as e:
                                logger.error(f"All attempts failed for relationships of type {rtype} between {src_lbl}->{tgt_lbl}: {e}")

                logger.info(f"Total attempted upload: {total_nodes} nodes and {total_relationships} relationships (note: fallbacks also run)")

                # --- Verification queries: log counts of relationship types
                try:
                    rel_counts = session.run("MATCH ()-[r]->() RETURN TYPE(r) AS rel_type, COUNT(r) AS cnt ORDER BY cnt DESC;")
                    rel_list = [(record["rel_type"], record["cnt"]) for record in rel_counts]
                    logger.info(f"Relationship counts by type: {rel_list}")
                except Exception as e:
                    logger.warning(f"Could not run verification rel counts: {e}")

                # Also log total counts
                try:
                    node_count = session.run("MATCH (n) RETURN COUNT(n) AS cnt").single()["cnt"]
                    rel_count = session.run("MATCH ()-[r]->() RETURN COUNT(r) AS cnt").single()["cnt"]
                    logger.info(f"Database totals after upload: nodes={node_count}, relationships={rel_count}")
                except Exception as e:
                    logger.warning(f"Could not retrieve db totals: {e}")

        except Exception as e:
            logger.error(f"Error uploading graph data: {e}")
            raise

    def query_graph(self, cypher_query: str) -> List[dict]:
        if not self.is_connected():
            return []

        try:
            with self.driver.session() as session:
                result = session.run(cypher_query)
                return [record.data() for record in result]
        except Exception as e:
            logger.error(f"Error executing query: {e}")
            return []

    def get_graph_stats(self) -> dict:
        if not self.is_connected():
            return {}

        try:
            with self.driver.session() as session:
                node_count = session.run("MATCH (n) RETURN COUNT(n) as count").single()["count"]
                rel_count = session.run("MATCH ()-[r]-() RETURN COUNT(r) as count").single()["count"]
                return {"nodes": node_count, "relationships": rel_count}
        except Exception as e:
            logger.error(f"Error getting graph stats: {e}")
            return {"nodes": 0, "relationships": 0}


class GraphRAGChatbot:
    """GraphRAG backend class using spaCy for graph extraction."""

    def __init__(self, spacy_model: str = "en_core_web_sm"):
        self.document_processor = DocumentProcessor()

        neo4j_uri = os.getenv("NEO4J_URI", "bolt://localhost:7687")
        neo4j_user = os.getenv("NEO4J_USER", "neo4j")
        neo4j_password = os.getenv("NEO4J_PASSWORD", "password")
        self.graph_db = Neo4jGraphDB(neo4j_uri, neo4j_user, neo4j_password)

        # Initialize spaCy extractor
        try:
            self.graph_extractor = SpacyGraphExtractor(model_name=spacy_model)
            logger.info("spaCy graph extractor initialized")
        except Exception as e:
            logger.error(f"Failed to initialize spaCy: {e}")
            self.graph_extractor = None

    def extract_and_upload_graph(self, text: str) -> Optional[dict]:
        """Extract graph from text using spaCy and upload to Neo4j."""
        if not self.graph_extractor:
            logger.warning("Graph extractor not available")
            return None

        try:
            # Extract graph using spaCy
            graph_data = self.graph_extractor.extract_graph(text)

            if graph_data and (graph_data['nodes'] or graph_data['relationships']):
                # Upload to Neo4j
                self.graph_db.upload_graph_data_from_spacy(graph_data)

                # Get updated stats
                stats = self.graph_db.get_graph_stats()

                return {
                    "nodes": len(graph_data['nodes']),
                    "relationships": len(graph_data['relationships']),
                    "stats": stats
                }
        except Exception as e:
            logger.error(f"Error extracting graph: {e}")

        return None

    def process_and_upload_file(self, file_path: str, use_docling: bool = True) -> Optional[dict]:
        """
        Full pipeline for one file: parse it with Docling (chunked),
        run spaCy graph extraction on every chunk, and upload everything
        to Neo4j.

        This is the function to call from app.py whenever a user uploads
        a document - it replaces calling process_document() and
        extract_and_upload_graph() separately for large files, since it
        processes the document chunk-by-chunk instead of as one giant
        block of text.
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        docling_supported = {'.pdf', '.docx', '.pptx', '.html', '.htm',
                              '.png', '.jpg', '.jpeg'}

        chunks: List[str]

        if use_docling and DOCLING_AVAILABLE and extension in docling_supported:
            try:
                chunks = self.document_processor.extract_with_docling_chunks(str(file_path))
            except Exception as e:
                logger.warning(f"Docling chunking failed, falling back to single-block text: {e}")
                text = self.document_processor.process_document(str(file_path), use_docling=False)
                chunks = [text]
        else:
            text = self.document_processor.process_document(str(file_path), use_docling=False)
            chunks = [text]

        total_nodes = 0
        total_relationships = 0
        for chunk_text in chunks:
            if not chunk_text or not chunk_text.strip():
                continue
            result = self.extract_and_upload_graph(chunk_text)
            if result:
                total_nodes += result["nodes"]
                total_relationships += result["relationships"]

        stats = self.graph_db.get_graph_stats()
        return {
            "chunks_processed": len(chunks),
            "nodes": total_nodes,
            "relationships": total_relationships,
            "stats": stats
        }

    def retrieve_context_from_graph(self, query: str) -> str:
        """Return a short text context pulled from Neo4j (case-insensitive search)."""
        if not self.graph_db.is_connected():
            return ""

        try:
            # Extract entities from query using spaCy
            if self.graph_extractor:
                doc = self.graph_extractor.nlp(query)
                query_entities = [ent.text.lower() for ent in doc.ents]  # Convert to lowercase

                if query_entities:
                    # Search for specific entities mentioned in query (case-insensitive)
                    # Use toLower() in Cypher to make comparison case-insensitive
                    entity_filter = " OR ".join([f"toLower(n.name) CONTAINS '{ent}'" for ent in query_entities[:5]])
                    cypher_query = f"""
                    MATCH (n)-[r]-(m)
                    WHERE {entity_filter}
                    RETURN DISTINCT n.name as source_name, TYPE(r) as rel_type, m.name as target_name
                    LIMIT 10
                    """
                else:
                    # Default query if no entities found
                    cypher_query = """
                    MATCH (n)-[r]-(m)
                    RETURN DISTINCT n.name as source_name, TYPE(r) as rel_type, m.name as target_name
                    LIMIT 10
                    """
            else:
                cypher_query = """
                MATCH (n)-[r]-(m)
                RETURN DISTINCT n.name as source_name, TYPE(r) as rel_type, m.name as target_name
                LIMIT 10
                """

            results = self.graph_db.query_graph(cypher_query)

            if results:
                context = "Knowledge Graph Context:\n"
                for result in results:
                    context += f"- {result.get('source_name', 'Unknown')} {result.get('rel_type', 'RELATED_TO')} {result.get('target_name', 'Unknown')}\n"
                return context
            return ""
        except Exception as e:
            logger.error(f"Error retrieving context: {e}")
            return ""
